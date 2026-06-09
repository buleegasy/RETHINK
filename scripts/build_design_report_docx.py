from __future__ import annotations

import json
import math
import re
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "青少年心理支持智能体作品设计报告.docx"
SURVEY_CSV = Path("/Users/chenhaoran/Documents/心理竞赛/survey_export_2026-06-06.csv")
RAG_DIR = ROOT / "rag-psy-cbt" / "data" / "cleaned"
FIG_DIR.mkdir(parents=True, exist_ok=True)

FONT_PATH = "/System/Library/Fonts/STHeiti Medium.ttc"
FONT_LIGHT_PATH = "/System/Library/Fonts/STHeiti Light.ttc"


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_PATH if bold else FONT_LIGHT_PATH, size)


PALETTE = {
    "ink": "#111827",
    "muted": "#4B5563",
    "blue": "#1F4E79",
    "blue2": "#EAF1F8",
    "green": "#2F5D50",
    "green2": "#EDF5F2",
    "amber": "#7A5A00",
    "amber2": "#F7F1DD",
    "red": "#8A1F1F",
    "red2": "#F7EAEA",
    "line": "#CBD5E1",
    "paper": "#FFFFFF",
}


def hex_to_rgb(h: str):
    h = h.lstrip("#")
    return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))


def wrap_text(draw, text, fnt, width):
    lines = []
    for para in str(text).split("\n"):
        if not para:
            lines.append("")
            continue
        line = ""
        for ch in para:
            trial = line + ch
            if draw.textbbox((0, 0), trial, font=fnt)[2] <= width:
                line = trial
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    return lines


def centered_text(draw, box, text, fnt, fill=PALETTE["ink"], line_gap=6):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, fnt, x2 - x1 - 24)
    line_h = fnt.size + line_gap
    total_h = len(lines) * line_h - line_gap
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=fnt, fill=hex_to_rgb(fill))
        y += line_h


def draw_round_rect(draw, box, fill, outline=PALETTE["line"], radius=6, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=hex_to_rgb(fill), outline=hex_to_rgb(outline), width=width)


def arrow(draw, start, end, fill=PALETTE["blue"], width=4):
    draw.line([start, end], fill=hex_to_rgb(fill), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 13
    pts = [
        end,
        (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6)),
        (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(pts, fill=hex_to_rgb(fill))


def canvas(title: str, subtitle: str = "", size=(1600, 900)):
    img = Image.new("RGB", size, hex_to_rgb(PALETTE["paper"]))
    d = ImageDraw.Draw(img)
    d.line((48, 86, size[0] - 48, 86), fill=hex_to_rgb(PALETTE["line"]), width=3)
    d.text((48, 24), title, font=font(32, True), fill=hex_to_rgb(PALETTE["ink"]))
    if subtitle:
        d.text((48, 60), subtitle, font=font(17), fill=hex_to_rgb(PALETTE["muted"]))
    return img, d


def save(img, name):
    path = FIG_DIR / name
    img.save(path, quality=95)
    return path


def read_survey():
    import csv

    if not SURVEY_CSV.exists():
        return None
    with SURVEY_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))

    def count_single(col):
        counts = {}
        for row in rows:
            val = (row.get(col) or "").strip()
            if val:
                counts[val] = counts.get(val, 0) + 1
        return counts

    def count_multi(col):
        counts = {}
        for row in rows:
            for val in (row.get(col) or "").split("|"):
                val = val.strip()
                if val:
                    counts[val] = counts.get(val, 0) + 1
        return counts

    comments = []
    for row in rows:
        c = (row.get("开放反馈") or "").strip()
        if c and c not in {"无", "没有", "🈚️"}:
            comments.append(c)
    return {
        "n": len(rows),
        "single": {c: count_single(c) for c in ["就读阶段", "AI经验", "付费意愿", "开场偏好", "隐私偏好"]},
        "multi": {c: count_multi(c) for c in ["情绪表情", "求助场景", "痛点", "期待人设"]},
        "comments": comments,
    }


def read_rag_reference_inventory():
    files = []
    refs = {}
    internal_rows = []
    if not RAG_DIR.exists():
        return files, [], internal_rows

    for fp in sorted(RAG_DIR.glob("*.jsonl")):
        count = 0
        for line in fp.read_text(encoding="utf-8").splitlines():
            if not line.strip():
                continue
            count += 1
            obj = json.loads(line)
            content = obj.get("content", "")
            match = re.search(r"来源[:：]\s*([^。\n]+)", content)
            source_label = match.group(1).strip(" -") if match else obj.get("summary", "")
            urls = obj.get("source_urls") or []
            if urls:
                for url in urls:
                    item = refs.setdefault(url, {
                        "source_labels": set(),
                        "source_types": set(),
                        "evidence": set(),
                        "files": set(),
                        "chunks": set(),
                    })
                    if source_label:
                        item["source_labels"].add(source_label)
                    if obj.get("source_type"):
                        item["source_types"].add(obj["source_type"])
                    evidence = (obj.get("metadata") or {}).get("evidence_level")
                    if evidence:
                        item["evidence"].add(evidence)
                    item["files"].add(fp.name)
                    if obj.get("chunk_id"):
                        item["chunks"].add(obj["chunk_id"])
            elif fp.name != "synthetic_case_chunks.jsonl":
                internal_rows.append([
                    fp.name,
                    obj.get("chunk_id", ""),
                    obj.get("source_type", ""),
                    obj.get("summary", ""),
                ])
        files.append([fp.name, str(count)])

    reference_rows = []
    for idx, (url, item) in enumerate(sorted(refs.items()), 1):
        labels = "；".join(sorted(item["source_labels"])) or "RAG 知识库来源"
        if len(labels) > 120:
            labels = labels[:117] + "..."
        kind = "；".join(sorted(item["source_types"]))
        evidence = "；".join(sorted(item["evidence"]))
        file_hint = "；".join(sorted(item["files"]))
        reference_rows.append([
            str(idx),
            labels,
            f"{kind}\n{evidence}\n涉及 {len(item['chunks'])} 个 chunk；{file_hint}",
            url,
        ])
    return files, reference_rows, internal_rows


def set_run_eastasia(run, font_name):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_style_eastasia(style, font_name):
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def flow_figure(name, title, nodes, edges, columns=4, subtitle=""):
    img, d = canvas(title, subtitle)
    w, h = img.size
    rows = math.ceil(len(nodes) / columns)
    box_w, box_h = 260, 104
    x_gap = (w - 2 * 90 - columns * box_w) / max(1, columns - 1)
    y_gap = (h - 180 - rows * box_h) / max(1, rows - 1) if rows > 1 else 0
    boxes = {}
    for i, (key, label, fill) in enumerate(nodes):
        r, c = divmod(i, columns)
        x1 = 90 + c * (box_w + x_gap)
        y1 = 170 + r * (box_h + y_gap)
        box = (x1, y1, x1 + box_w, y1 + box_h)
        boxes[key] = box
        draw_round_rect(d, box, fill, radius=22)
        centered_text(d, box, label, font(23, True))
    for a, b in edges:
        ax1, ay1, ax2, ay2 = boxes[a]
        bx1, by1, bx2, by2 = boxes[b]
        start = (ax2, (ay1 + ay2) / 2) if bx1 > ax2 else ((ax1 + ax2) / 2, ay2)
        end = (bx1, (by1 + by2) / 2) if bx1 > ax2 else ((bx1 + bx2) / 2, by1)
        arrow(d, start, end)
    return save(img, name)


def table_figure(name, title, headers, rows, col_widths=None, subtitle=""):
    img, d = canvas(title, subtitle, size=(1600, 950))
    x0, y0, width = 72, 150, 1456
    if col_widths is None:
        col_widths = [width / len(headers)] * len(headers)
    else:
        scale = width / sum(col_widths)
        col_widths = [v * scale for v in col_widths]
    row_h = max(56, min(98, (img.size[1] - y0 - 60) / (len(rows) + 1)))
    y = y0
    x = x0
    for j, head in enumerate(headers):
        d.rectangle((x, y, x + col_widths[j], y + row_h), fill=hex_to_rgb(PALETTE["blue2"]), outline=hex_to_rgb(PALETTE["line"]), width=2)
        centered_text(d, (x + 6, y + 4, x + col_widths[j] - 6, y + row_h - 4), head, font(21, True))
        x += col_widths[j]
    y += row_h
    for i, row in enumerate(rows):
        x = x0
        fill = "#FFFFFF" if i % 2 == 0 else "#F7FAF8"
        for j, cell in enumerate(row):
            d.rectangle((x, y, x + col_widths[j], y + row_h), fill=hex_to_rgb(fill), outline=hex_to_rgb(PALETTE["line"]), width=2)
            centered_text(d, (x + 8, y + 4, x + col_widths[j] - 8, y + row_h - 4), str(cell), font(18))
            x += col_widths[j]
        y += row_h
    return save(img, name)


def bar_figure(name, title, labels, values, subtitle=""):
    img, d = canvas(title, subtitle)
    x0, y0, x1, y1 = 130, 190, 1480, 760
    d.line((x0, y1, x1, y1), fill=hex_to_rgb(PALETTE["line"]), width=3)
    d.line((x0, y0, x0, y1), fill=hex_to_rgb(PALETTE["line"]), width=3)
    max_v = max(values)
    bar_w = (x1 - x0) / len(values) * 0.62
    gap = (x1 - x0) / len(values)
    colors = [PALETTE["blue"], PALETTE["green"], PALETTE["amber"], PALETTE["red"]]
    for i, (lab, val) in enumerate(zip(labels, values)):
        cx = x0 + gap * i + gap / 2
        bh = (y1 - y0 - 30) * val / max_v
        d.rounded_rectangle((cx - bar_w / 2, y1 - bh, cx + bar_w / 2, y1), radius=10, fill=hex_to_rgb(colors[i % 4]))
        txt = str(val)
        bbox = d.textbbox((0, 0), txt, font=font(23, True))
        d.text((cx - (bbox[2] - bbox[0]) / 2, y1 - bh - 36), txt, font=font(23, True), fill=hex_to_rgb(PALETTE["ink"]))
        for k, line in enumerate(wrap_text(d, lab, font(17), gap - 18)[:2]):
            bbox = d.textbbox((0, 0), line, font=font(17))
            d.text((cx - (bbox[2] - bbox[0]) / 2, y1 + 22 + k * 23), line, font=font(17), fill=hex_to_rgb(PALETTE["muted"]))
    return save(img, name)


def pie_figure(name, title, labels, values, subtitle=""):
    img, d = canvas(title, subtitle)
    cx, cy, r = 520, 500, 260
    total = sum(values)
    start = -90
    colors = [PALETTE["blue"], PALETTE["green"], PALETTE["amber"], PALETTE["red"], "#6E5C9A", "#4F7A8A"]
    for lab, val, col in zip(labels, values, colors):
        end = start + 360 * val / total
        d.pieslice((cx - r, cy - r, cx + r, cy + r), start, end, fill=hex_to_rgb(col), outline="white", width=4)
        start = end
    lx, ly = 940, 260
    for i, (lab, val, col) in enumerate(zip(labels, values, colors)):
        y = ly + i * 70
        d.rounded_rectangle((lx, y, lx + 42, y + 42), radius=8, fill=hex_to_rgb(col))
        d.text((lx + 58, y + 4), f"{lab}：{val}分", font=font(25), fill=hex_to_rgb(PALETTE["ink"]))
    return save(img, name)


def timeline_figure(name, title, items, subtitle=""):
    img, d = canvas(title, subtitle)
    x0, x1, y = 150, 1450, 450
    d.line((x0, y, x1, y), fill=hex_to_rgb(PALETTE["blue"]), width=6)
    gap = (x1 - x0) / (len(items) - 1)
    for i, (label, detail) in enumerate(items):
        x = x0 + i * gap
        d.ellipse((x - 18, y - 18, x + 18, y + 18), fill=hex_to_rgb(PALETTE["green"]), outline="white", width=4)
        box = (x - 112, y + (42 if i % 2 == 0 else -180), x + 112, y + (166 if i % 2 == 0 else -56))
        draw_round_rect(d, box, PALETTE["green2"], radius=18)
        centered_text(d, box, f"{label}\n{detail}", font(18), line_gap=3)
        if i % 2 == 0:
            d.line((x, y + 18, x, y + 42), fill=hex_to_rgb(PALETTE["line"]), width=3)
        else:
            d.line((x, y - 18, x, y - 56), fill=hex_to_rgb(PALETTE["line"]), width=3)
    return save(img, name)


def radar_figure(name, title, labels, values, subtitle=""):
    img, d = canvas(title, subtitle)
    cx, cy, r = 800, 500, 280
    n = len(labels)
    for level in range(1, 6):
        pts = []
        rr = r * level / 5
        for i in range(n):
            ang = -math.pi / 2 + 2 * math.pi * i / n
            pts.append((cx + rr * math.cos(ang), cy + rr * math.sin(ang)))
        d.line(pts + [pts[0]], fill=hex_to_rgb(PALETTE["line"]), width=2)
    pts = []
    for i, val in enumerate(values):
        ang = -math.pi / 2 + 2 * math.pi * i / n
        pts.append((cx + r * val / 5 * math.cos(ang), cy + r * val / 5 * math.sin(ang)))
        lx = cx + (r + 95) * math.cos(ang)
        ly = cy + (r + 58) * math.sin(ang)
        centered_text(d, (lx - 95, ly - 30, lx + 95, ly + 30), labels[i], font(19), fill=PALETTE["ink"])
    d.polygon(pts, fill=(39, 107, 122, 90), outline=hex_to_rgb(PALETTE["blue"]))
    d.line(pts + [pts[0]], fill=hex_to_rgb(PALETTE["blue"]), width=5)
    return save(img, name)


def create_figures():
    survey = read_survey()
    figs = []
    figs.append(flow_figure("figure01_project_position.png", "Figure 1  项目定位：低门槛心理支持入口", [
        ("a", "青少年倾诉", PALETTE["blue2"]), ("b", "大模型理解", PALETTE["green2"]),
        ("c", "ABC事实剥离", PALETTE["amber2"]), ("d", "RAG安全检索", PALETTE["blue2"]),
        ("e", "个性化回应", PALETTE["green2"]), ("f", "现实求助路径", PALETTE["amber2"]),
    ], [("a", "b"), ("b", "c"), ("c", "d"), ("d", "e"), ("e", "f")], columns=3))
    figs.append(table_figure("figure02_problem_solution.png", "Figure 2  问题识别与解决路径", ["发现的问题", "解决设计", "验证方式"], [
        ["事实与评价混杂", "ABC事实剥离与监控摄像头标准", "样例对话逐条校验"],
        ["普通聊天容易过度安慰", "共情 + 证据检验 + 小行动", "对比普通回应与结构化回应"],
        ["危机场景不能普通劝说", "安全优先路由覆盖普通CBT", "危机测试用例检查"],
        ["知识来源不透明", "RAG来源索引和字段规范", "验证脚本检查元数据"],
    ], [2.2, 2.3, 2.0]))
    figs.append(flow_figure("figure02a_questionnaire.png", "Figure 2A  开题调查问卷如何转化设计决策", [
        ("a", "匿名专项调查", PALETTE["blue2"]),
        ("b", "就读阶段 / AI 使用经历", PALETTE["green2"]),
        ("c", "负面情绪表情包偏好", PALETTE["amber2"]),
        ("d", "求助场景与反感痛点", PALETTE["blue2"]),
        ("e", "记忆边界 / 角色偏好 / 开场界面", PALETTE["green2"]),
        ("f", "反推产品架构与话术设计", PALETTE["amber2"]),
    ], [("a", "b"), ("b", "c"), ("c", "d"), ("d", "e"), ("e", "f")], columns=3, subtitle="问卷的作用不是收集“标准答案”，而是把真实偏好转成产品设计约束。"))
    figs.append(table_figure("figure02b_questionnaire_items.png", "Figure 2B  问卷 10 题与设计决策映射", ["题号", "问什么", "为什么问", "怎么用到设计里"], [
        ["1", "就读阶段", "区分初中/高中/更高学段的表达差异", "决定语气、词汇和案例难度"],
        ["2", "AI 倾诉经验", "判断熟悉度与使用门槛", "决定开场是否需要解释功能"],
        ["3", "负面情绪表情包", "观察低谷状态下的直觉表达", "帮助设计低输入压力入口"],
        ["4", "更愿意求助 AI 的场景", "定位真实使用时机", "确定重点场景：深夜、考试、人际、隐私"],
        ["5", "最反感的痛点", "找出劝退原因", "规避鸡汤、盘问、模板化、隐私焦虑"],
        ["6", "记忆边界偏好", "确认用户对记忆的控制感", "设计一键清空/阅后即焚等边界说明"],
        ["7", "希望 AI 扮演的角色", "区分树洞、分析师、导师、医生、乐子人", "决定产品人格和模式切换"],
        ["8", "付费意愿", "判断价值感与接受度", "辅助后续产品化判断"],
        ["9", "哪种开场界面更卸防", "比较表情包/温和寒暄/压力评估", "决定首屏形态"],
        ["10", "开放建议", "收集用户原声吐槽和期待", "补充产品迭代线索"],
    ], [0.5, 1.3, 1.9, 2.8], subtitle="这份问卷不是只看“喜不喜欢”，而是把交互方式、角色人格、隐私控制和首屏形态拆成可设计的变量。"))
    if survey:
        pain = survey["multi"]["痛点"]
        pain_labels = ["模板化", "空洞鸡汤", "无记忆", "量表盘问", "隐私焦虑"]
        pain_keys = ["stiff_template", "toxic_positivity", "no_memory", "clinical_survey", "privacy_paranoia"]
        figs.append(bar_figure("figure02c_survey_painpoints.png", "Figure 2C  问卷痛点分布（n=71）", pain_labels, [pain.get(k, 0) for k in pain_keys], subtitle="痛点用于解释为什么系统不能只做安慰型聊天，而要做结构化、可控、非模板化回应。"))
        scene = survey["multi"]["求助场景"]
        scene_labels = ["高度隐私", "深夜emo", "社交疲惫", "人际纠纷", "空虚无聊", "情感困扰", "考试失利"]
        scene_keys = ["highly_private", "midnight_emo", "social_fatigue", "relationship_dispute", "existential_void", "puppy_love", "academic_crash"]
        figs.append(bar_figure("figure02d_survey_scenarios.png", "Figure 2D  AI心理求助场景分布（n=71）", scene_labels, [scene.get(k, 0) for k in scene_keys], subtitle="高隐私、深夜情绪和社交疲惫是智能体低门槛入口的重要依据。"))
        opening = survey["single"]["开场偏好"]
        figs.append(bar_figure("figure02e_opening_preference.png", "Figure 2E  首屏开场偏好分布（n=71）", ["表情包选择", "温和寒暄", "压力评估"], [opening.get("emoji_selector", 0), opening.get("soft_chitchat", 0), opening.get("clinical_scale", 0)], subtitle="表情包与温和寒暄明显高于量表式开场，支持低压力进入设计。"))
    else:
        figs.append(table_figure("figure02c_survey_painpoints.png", "Figure 2C  问卷痛点分布", ["状态"], [["未找到问卷CSV，待补充数据图"]]))
        figs.append(table_figure("figure02d_survey_scenarios.png", "Figure 2D  AI心理求助场景分布", ["状态"], [["未找到问卷CSV，待补充数据图"]]))
        figs.append(table_figure("figure02e_opening_preference.png", "Figure 2E  首屏开场偏好分布", ["状态"], [["未找到问卷CSV，待补充数据图"]]))
    figs.append(flow_figure("figure03_system_architecture.png", "Figure 3  系统总体架构", [
        ("u", "用户输入", PALETTE["blue2"]), ("n", "自然语言理解", PALETTE["green2"]), ("a", "ABC模块", PALETTE["amber2"]), ("r", "风险路由", PALETTE["red2"]),
        ("k", "RAG知识库", PALETTE["blue2"]), ("g", "回应生成", PALETTE["green2"]), ("o", "输出与记录", PALETTE["amber2"]),
    ], [("u", "n"), ("n", "a"), ("n", "r"), ("a", "k"), ("r", "k"), ("k", "g"), ("g", "o")], columns=4))
    figs.append(flow_figure("figure04_abc_flow.png", "Figure 4  ABC事实剥离工作流", [
        ("s", "扫描情绪词/绝对词", PALETTE["blue2"]), ("f", "剔除主观滤镜", PALETTE["green2"]),
        ("a", "提取客观事实A", PALETTE["amber2"]), ("b", "标注信念解释B", PALETTE["blue2"]),
        ("c", "整理情绪行为C", PALETTE["green2"]), ("q", "摄像头标准校验", PALETTE["amber2"]),
    ], [("s", "f"), ("f", "a"), ("a", "b"), ("b", "c"), ("c", "q")], columns=3))
    figs.append(table_figure("figure05_cognitive_distortions.png", "Figure 5  青少年高发认知偏差矩阵", ["偏差", "典型表达", "解决策略"], [
        ["非黑即白", "没做到最好就是失败", "保留中间状态和局部改进"],
        ["读心术", "他肯定讨厌我", "区分事实与动机推测"],
        ["灾难化", "我以后完了", "降低确定性预测"],
        ["过度概括", "一次失误=一直失败", "限定事件范围"],
        ["个人化", "老师脸色差一定因为我", "检查因果证据"],
    ], [1.2, 2.6, 2.5]))
    figs.append(flow_figure("figure06_risk_route.png", "Figure 6  安全优先风险路由", [
        ("i", "输入风险扫描", PALETTE["blue2"]), ("c", "危机风险", PALETTE["red2"]),
        ("e", "升高风险", PALETTE["amber2"]), ("n", "普通支持", PALETTE["green2"]),
        ("s", "现实支持/转介", PALETTE["red2"]), ("p", "澄清与资源建议", PALETTE["amber2"]), ("b", "CBT事实剥离", PALETTE["green2"]),
    ], [("i", "c"), ("i", "e"), ("i", "n"), ("c", "s"), ("e", "p"), ("n", "b")], columns=4))
    figs.append(table_figure("figure07_rag_layers.png", "Figure 7  RAG知识来源分层", ["层级", "来源", "用途"], [
        ["权威指南", "WHO / NICE / AAP / AACAP", "风险识别与安全边界"],
        ["中国政策", "教育部心理健康文件", "校园应用与家校协同"],
        ["CBT理论", "Beck Institute等", "事实剥离与认知重评"],
        ["合成示例", "fictional_cross_checked", "对话格式演示"],
    ], [1.2, 2.4, 2.8]))
    figs.append(bar_figure("figure08_data_scale.png", "Figure 8  清洗后知识材料规模", ["CBT", "权威指南", "专业理论", "对话", "政策", "合成场景", "来源索引", "RAG规则"], [6, 28, 40, 4, 2, 120, 81, 41]))
    figs.append(timeline_figure("figure09_process_timeline.png", "Figure 9  项目创作与解决问题过程", [
        ("需求拆解", "对齐参赛方向"), ("问题发现", "识别AI心理风险"), ("规则设计", "ABC与安全边界"),
        ("知识整理", "RAG来源分层"), ("原型测试", "破冰开场验证"), ("报告成稿", "图文证据链"),
    ]))
    figs.append(flow_figure("figure09a_emoji_entry.png", "Figure 9A  情绪表情包低压力入口机制", [
        ("a", "用户不想打长句", PALETTE["blue2"]),
        ("b", "选择1-3个表情", PALETTE["green2"]),
        ("c", "映射情绪倾向", PALETTE["amber2"]),
        ("d", "降低开场压力", PALETTE["blue2"]),
        ("e", "进入文字或陪伴模式", PALETTE["green2"]),
    ], [("a", "b"), ("b", "c"), ("c", "d"), ("d", "e")], columns=3, subtitle="问卷中表情包题不是装饰，而是把复杂情绪转化为低负担输入。"))
    figs.append(flow_figure("figure09b_sentence_compression.png", "Figure 9B  长句到短句的表达压缩流程", [
        ("raw", "用户长句/混乱倾诉", PALETTE["blue2"]),
        ("split", "拆分事实/情绪/诉求", PALETTE["green2"]),
        ("rank", "保留最高优先级信息", PALETTE["amber2"]),
        ("short", "生成短句开场", PALETTE["blue2"]),
        ("choice", "给用户选择权", PALETTE["green2"]),
    ], [("raw", "split"), ("split", "rank"), ("rank", "short"), ("short", "choice")], columns=3, subtitle="目标不是删减信息，而是把第一轮交互压到青少年愿意读、愿意回的长度。"))
    figs.append(table_figure("figure09c_prompt_engineering.png", "Figure 9C  拟人化回应提示词工程框架", ["提示词约束", "解决的问题", "输出要求"], [
        ["先共情再分析", "避免上来讲道理或盘问", "先命名情绪，再进入事实拆分"],
        ["短句优先", "避免大段说教", "每轮回应控制信息密度"],
        ["少术语", "降低心理咨询距离感", "用学生能听懂的话解释"],
        ["不模板化", "回应“不是完全套话”的反馈", "结合用户原话和场景生成"],
        ["给选择权", "降低被控制感", "提供倾听/分析/小行动选项"],
        ["小行动", "避免建议空泛", "给当下可做的一步"],
    ], [1.5, 2.2, 2.5], subtitle="提示词工程不是让AI假装像人，而是约束它少说套话、保留温度、能具体帮忙。"))
    figs.append(flow_figure("figure10_opening_test.png", "Figure 10  首次聊天破冰测试流程", [
        ("intro", "进入说明页", PALETTE["blue2"]), ("card", "随机话术卡片", PALETTE["green2"]),
        ("choice", "继续/跳过/退出", PALETTE["amber2"]), ("record", "记录反应", PALETTE["blue2"]),
        ("result", "生成结果概览", PALETTE["green2"]),
    ], [("intro", "card"), ("card", "choice"), ("choice", "record"), ("record", "result")], columns=3))
    figs.append(radar_figure("figure11_scene_radar.png", "Figure 11  青少年场景覆盖示意", ["考试压力", "同伴关系", "师生互动", "亲子沟通", "自我评价", "危机识别"], [5, 4, 4, 3, 5, 4]))
    figs.append(pie_figure("figure12_scoring_pie.png", "Figure 12  评分维度占比", ["知识掌握", "目标完成度", "专业质量", "创新突破", "心理学运用", "心理素养应用"], [15, 20, 15, 20, 10, 20]))
    figs.append(table_figure("figure13_response_template.png", "Figure 13  个性化回应模板", ["步骤", "输出重点", "目的"], [
        ["1", "情绪接纳", "让用户感觉被理解"],
        ["2", "事实剥离", "降低情绪化推断"],
        ["3", "偏差识别", "看见思维模式"],
        ["4", "平衡解释", "形成更可承受的理解"],
        ["5", "小行动", "把支持落到具体步骤"],
    ], [0.8, 2.4, 3.0]))
    figs.append(flow_figure("figure14_validation_chain.png", "Figure 14  可验证证据链", [
        ("rules", "规则文档", PALETTE["blue2"]), ("data", "知识库JSONL", PALETTE["green2"]),
        ("script", "验证脚本", PALETTE["amber2"]), ("html", "测试页面", PALETTE["blue2"]),
        ("shot", "过程截图", PALETTE["green2"]), ("pdf", "PDF提交", PALETTE["amber2"]),
    ], [("rules", "data"), ("data", "script"), ("script", "html"), ("html", "shot"), ("shot", "pdf")], columns=3))
    figs.append(table_figure("figure15_ethics_guardrails.png", "Figure 15  伦理与安全护栏", ["允许", "禁止", "原因"], [
        ["心理教育、情绪支持", "疾病诊断", "避免替代专业评估"],
        ["结构化整理、小行动", "药物处方", "避免医学风险"],
        ["现实求助建议", "伤害方法", "保护未成年人安全"],
        ["权威资料引用", "伪造病例", "保证来源可核验"],
    ], [2.0, 2.0, 2.5]))
    figs.append(table_figure("figure16_submission_package.png", "Figure 16  参赛提交材料包", ["文件", "内容", "格式检查"], [
        ["作品设计报告", "背景、目的、功能、创新点、支撑数据", "PDF，查重≤30%"],
        ["过程性证明材料", "流程截图、成果视频链接", "正文图片，链接可播放"],
        ["作品查重报告", "知网/维普/万方等检测", "PDF，供组委会核验"],
    ], [1.5, 3.2, 1.8]))
    return figs


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Calibri"
                set_run_eastasia(run, "黑体")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    set_style_eastasia(normal, "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        st = styles[name]
        st.font.name = "Calibri"
        set_style_eastasia(st, "黑体")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 16)
        st.paragraph_format.space_after = Pt(6)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.3))
    add_caption(doc, caption)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def breakable_url(url: str):
    for ch in ["/", "?", "&", "=", "-", "_", "."]:
        url = url.replace(ch, ch + "\u200b")
    return url


def add_reference_list(doc, rows):
    for idx, label, usage, url in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.05
        head = p.add_run(f"[{idx}] ")
        head.bold = True
        head.font.size = Pt(9)
        body = p.add_run(f"{label}。来源类型与用途：{usage.replace(chr(10), '；')}。URL：{breakable_url(url)}")
        body.font.size = Pt(8.5)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")
    p.add_run("\n" + body)
    doc.add_paragraph()


def build_doc():
    figs = create_figures()
    survey = read_survey()
    n = survey["n"] if survey else 0
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("青少年心理支持智能体作品设计报告")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    set_run_eastasia(r, "黑体")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("基于 CBT 事实剥离与 RAG 安全边界的个性化心理支持原型")
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor.from_string("555555")
    doc.add_paragraph("参赛方向：智能体应用方向    适用对象：初中生、高中生及校园心理健康教育场景").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("版本：Word 草案 V0.2").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(
        "本项目面向青少年在学业压力、同伴关系、自我评价、亲子沟通等高频心理场景中的情绪支持需求，"
        "设计一套基于主流人工智能大模型的心理主题智能体原型。系统以自然语言处理为入口，识别用户倾诉中的事实、"
        "信念、情绪、行为倾向与风险信号；以 CBT 认知行为疗法中的 ABC 模型和自动化思维记录为核心解释框架，"
        "帮助用户区分“客观发生的事件”和“脑中出现的解释”；以权威指南与政策资料构建 RAG 知识库，"
        "确保系统在危机、升高风险、普通支持等不同场景中执行分层回应。"
    )
    add_callout(
        doc,
        "项目定位",
        "本项目不进行心理诊断，不替代心理咨询、精神科评估、医学检查或危机急救；其定位是心理教育、情绪支持、结构化整理、自助练习与现实求助路径建议。"
    )

    doc.add_heading("一、项目背景与问题发现", level=1)
    doc.add_heading("1.1 青少年心理需求", level=2)
    doc.add_paragraph(
        "青少年阶段处于自我概念发展、学业竞争增强、人际归属需求上升和家庭关系重塑的交汇期。"
        "许多心理困扰并不一定以正式疾病形式出现，而是表现为考试失利后的自我否定、同伴互动中的过度推测、"
        "教师反馈后的羞耻感，或家长期待下的压力累积。"
    )
    doc.add_heading("1.2 项目开题调查问卷", level=2)
    doc.add_paragraph(
        "在项目刚开始时，我们没有直接进入功能开发，而是先做了一轮关于心理 AI 交互偏好的匿名专项调查。"
        "这套问卷对应的实际页面是 `survey.rethink.buleegasy.space`，配套还有一个数据大屏 `re-think-dashboard.pages.dev` 用来汇总样本和偏好分布。"
        f"截至本次导出的 `survey_export_2026-06-06.csv`，问卷共记录 {n} 份提交。"
        "在项目早期，大屏的角色更像一个需求看板：它用来观察样本逐步累积后，各类偏好会怎样分布，而不是一开始就拿来做结论展示。"
        "问卷不是单纯问“你喜不喜欢 AI”，而是围绕中学生在心理健康辅导智能体中的交互响应机制、回复话术风格和情感回馈模式展开。"
        "它把一个看似抽象的课题拆成十个可回答的问题：受访者是什么学段、以前有没有用 AI 倾诉、负面状态下最像哪种表情包、"
        "在什么场景下会更愿意找 AI 而不是家长或老师、最反感 AI 心理对话里的哪些痛点、希望 AI 记住多少内容、"
        "更愿意 AI 扮演什么角色、对付费大致持什么态度、哪一种首屏开场最能卸下防线，以及还有什么真实建议和吐槽。"
    )
    doc.add_paragraph(
        "问卷的价值在于把抽象需求变成可执行的设计约束。第一，它让我们确认了用户并不喜欢一上来就被“盘问”，"
        "而是希望系统先给出低压力入口；第二，它把“鸡汤式鼓励”“测评式盘问”“机械模板”和“隐私顾虑”这些痛点具体化，"
        "使我们明确必须采用先共情、再剥离事实、最后再给选择权的流程；第三，它帮助我们看见不同用户对记忆边界的差异，"
        "有人希望像老朋友一样记住过往，有人则明确需要一键清空或阅后即焚，所以系统不能只做单一路径，而要把记忆控制权写给用户。"
        "此外，问卷对“树洞 / 分析师 / 严厉导师 / 心理医生 / 乐子人”等角色偏好，也让我们意识到产品人格不能过于单一，"
        "首屏必须支持不同类型的进入方式，既能温和陪伴，也能给出结构化分析。"
    )
    doc.add_paragraph(
        "更重要的是，这份问卷和数据大屏不是做完就结束，而是直接进入项目早期的迭代循环：我们先通过问卷确认青少年的真实偏好，"
        "再把这些偏好翻译成产品约束，最后回到首屏开场、模式切换、隐私边界和回应风格上做原型验证。"
        "换句话说，问卷不是附属材料，而是项目最早的一次设计校准。"
    )
    add_figure(doc, figs[3], "Figure 1. 问卷题项到模块设计的追踪矩阵。")
    doc.add_paragraph(
        "开放反馈进一步说明了为什么这些分布会转化为具体模块。用户不是简单要求“更温柔”，而是在反复表达三个更具体的需求："
        "第一，AI 不能只安慰，还要能分析；第二，AI 不能像模板，要能记住上下文和个人经历；第三，AI 给出的建议要落到当下可做的小动作。"
        "这些原话直接成为后续模块设计的动机。"
    )
    doc.add_paragraph(
        "其中两条较长的开放评论尤其能说明项目从“普通聊天工具”转向“问卷驱动的心理支持智能体”的原因。"
        "#survey-mq2... 在 2026/6/6 21:19:45 写道：“我希望他/她能像我的好朋友一样能鼓励我共情我，不要很生硬，就像聊天一样。”"
        "这句话直接推动了我们对低压力开场、同伴式语气和拟人化提示词的设计：系统不能像冷冰冰的问诊表，也不能只输出标准化心理教育，"
        "而要先让用户感觉自己是在和一个能接住情绪的人说话。"
    )
    doc.add_paragraph(
        "另一条 #survey-mq1... 在 2026/6/6 10:41:32 写道：“嗯，deepseek（目前只和dd聊过）能很理性的分析，"
        "又能用很温暖的句子戳中人心，但是建议有些空泛，比如抑郁发作时什么都不想动情绪很糟糕不知道做什么能缓解，"
        "说实话心理医生的建议也有限，我把希望寄托在dd上，但是都是大差不差的回答 希望可以更人性化一点，因为用久了会发现，"
        "其实回答有一些模板化，然后希望可以更专业一点的评估，虽然不能代替正规的确诊，但是也许可以反映一些当下的情况，"
        "希望可以让更多的人提早发现自己的心理问题，因为我自己确诊前真的就是过完发作期（非常惨烈的大哭）之后大脑应该因为自我保护选择模糊这段记忆，"
        "就是一种极端的空虚与茫然，不知道为什么而哭，然后就会自然而然的认为自己现在目前状态还可以，就不会想到去医院之类的事情，"
        "就这样反复会让自己在大多数时间里觉得自己是正常的，而过了发作期，也就不会想要去医院，所以真的是这样反复拖了很久，"
        "导致之后很严重了才去开药，所以希望可以让大家早点发现早点确诊早点药物治疗不要像我一样。”"
        "这条反馈暴露出更复杂的需求：用户既需要温暖和人性化，又需要比空泛建议更具体的行动提示；"
        "既期待专业评估感，又明确不能让 AI 替代正规诊断。因此，项目后续把“短句化回应、微小行动建议、RAG 专业知识支撑、"
        "非诊断边界和危机分流”放在同一条设计链路中，而不是只做语气优化。"
    )
    add_table(doc, ["开放反馈原话", "暴露的问题", "对应模块"], [
        ["“希望不要再只是安慰我 合理去分析我的问题”", "单纯共情不足，用户需要结构化分析。", "ABC事实剥离、认知偏差识别"],
        ["“不是那种完全套话”", "模板化回应降低信任。", "个性化回应生成、上下文引用"],
        ["“不要太套路化了，感觉每次都能猜到ai说什么”", "固定话术会让心理支持失真。", "RAG检索与场景化生成"],
        ["“辅导时可以多给一点当下能做的一些微小能缓解压力行为的建议”", "建议必须可执行、低成本。", "小行动建议模块"],
        ["“记住我之前说过的事”", "用户需要连续性支持。", "用户可控记忆与上下文管理"],
        ["“可以理性但不能太理性，要感性”", "理性分析与情绪承接需要平衡。", "情感计算与回应风格调节"],
        ["“对于一些内向的同学，他们可以毫无顾虑的接受ai心理辅导”", "匿名、低压力入口具有价值。", "隐私边界、低压力首屏"],
    ], [2.6, 2.0, 1.8])
    doc.add_heading("1.3 我们发现的关键问题", level=2)
    doc.add_paragraph(
        "在梳理项目需求和早期样例时，我们发现，如果只让大模型“安慰用户”，容易出现三类问题："
        "第一，模型可能把用户的主观推断当作事实；第二，模型可能输出空泛鼓励，不能帮助用户形成具体行动；"
        "第三，危机场景下如果继续做普通认知重构，会延误现实求助。因此，项目没有把目标设定为“做一个会聊天的机器人”，"
        "而是把目标设定为“做一个能按心理安全规则工作、能被验证的支持型智能体”。"
    )

    doc.add_heading("二、项目目的", level=1)
    doc.add_paragraph(
        "本项目目标不是单纯完成一个聊天界面，而是围绕“青少年是否愿意使用、系统是否能安全回应、回应是否具备心理学结构、方案是否可验证”"
        "四个问题建立完整原型。为此，项目将用户研究、心理学规则、知识库治理和原型测试放在同一条设计链路中："
        "先通过问卷识别青少年对心理 AI 的真实偏好，再把偏好转化为交互原则；随后用 CBT 和风险路由约束智能体回应，"
        "最后通过测试用例、知识库脚本和过程性材料证明系统可以被复核。"
    )
    doc.add_paragraph("具体目标包括：")
    add_numbered(doc, [
        "识别青少年倾诉中的客观事实、主观解释、情绪与行为倾向。",
        "根据 CBT ABC 模型帮助用户区分“发生了什么”和“我如何理解它”。",
        "识别非黑即白、读心术、灾难化、过度概括、个人化等常见认知偏差。",
        "根据风险等级调用不同响应策略，普通压力场景提供情绪支持和小行动建议，危机场景优先给出现实求助路径。",
        "通过 RAG 知识库引入权威指南、政策资料和 CBT 理论，提升输出的可解释性和可追溯性。",
    ])
    doc.add_heading("2.1 研究问题", level=2)
    add_table(doc, ["研究问题", "对应设计任务", "验证材料"], [
        ["RQ1：青少年在什么心理场景下愿意使用 AI？", "开题问卷识别高频求助场景、反感痛点和角色偏好。", "问卷页面、数据大屏、问卷题项映射图。"],
        ["RQ2：智能体如何避免把主观评价当成事实？", "建立 ABC 事实剥离 SOP 和监控摄像头标准。", "核心规则文档、样例对话、功能测试用例。"],
        ["RQ3：系统如何处理普通压力与危机风险的差异？", "设计风险路由和 crisis 知识块优先召回机制。", "RAG 质量规则、危机测试用例、伦理边界图。"],
        ["RQ4：知识来源如何保证可追溯？", "区分权威指南、政策文件、CBT 理论和合成示例。", "来源索引、JSONL 元数据、验证脚本。"],
    ], [1.8, 2.5, 2.0])
    doc.add_heading("2.2 设计原则", level=2)
    doc.add_paragraph(
        "基于问卷反馈和心理安全要求，项目形成四条设计原则。第一，低压力进入：用户不需要一开始就完整讲述问题，"
        "可以通过表情包、短句或模式选择进入对话。第二，先接纳再分析：系统先确认情绪，再拆分事实和解释，避免上来盘问或说教。"
        "第三，用户可控：对记忆、隐私、对话节奏和是否继续分析保留选择权。第四，安全优先：当出现危机信号时，"
        "系统停止普通陪聊和认知重构，转向现实求助和专业资源。"
    )

    doc.add_heading("三、系统设计与功能", level=1)
    doc.add_heading("3.1 总体架构", level=2)
    doc.add_paragraph(
        "总体架构的设计动机来自问卷中的两个集中反馈：一类用户担心 AI “说话机械套路”，另一类用户希望 AI “合理去分析我的问题”。"
        "因此，系统不能只依赖单次生成，而需要把理解、分析、安全、检索和输出拆成可解释的模块。"
        "系统由自然语言理解、ABC 事实剥离、情绪识别、风险路由、RAG 知识检索和回应生成六个模块组成。"
        "用户输入首先进入自然语言理解模块，系统提取情绪词、绝对化表达、他人动机推测和风险信号；"
        "随后通过 ABC 模块把事实、信念和情绪行为分开；风险路由模块决定是否进入普通支持、升高风险支持或危机安全处置。"
    )
    doc.add_paragraph(
        "在工程实现上，系统采用“输入解析-安全判断-知识检索-回应生成-输出校验”的流水线结构。这样做的原因是，"
        "心理支持场景中的错误往往不是单一文本生成错误，而是由多个环节叠加造成：如果输入解析阶段没有识别出危机词，"
        "后续回应可能错误进入普通安慰；如果知识检索阶段没有召回权威指南，系统可能输出没有边界的建议；"
        "如果输出校验阶段没有检查诊断化表达，回应可能超出非诊断定位。因此，本项目把安全检查前置，并在回应生成后再次进行边界校验。"
    )
    add_table(doc, ["模块", "输入", "处理逻辑", "输出"], [
        ["自然语言理解", "用户原始倾诉", "抽取事实词、情绪词、绝对化词、风险词。", "结构化输入特征"],
        ["ABC事实剥离", "结构化输入特征", "按监控摄像头标准区分 A/B/C。", "事实、信念、情绪行为"],
        ["风险路由", "情绪强度与风险信号", "判断 normal/elevated/crisis。", "回应模式与召回优先级"],
        ["RAG检索", "场景、技术、风险等级", "按证据层级召回指南、政策、CBT理论。", "可追溯知识片段"],
        ["回应生成", "ABC结果与知识片段", "共情、重评、小行动或现实求助。", "最终支持回应"],
        ["输出校验", "草稿回应", "检查诊断化、过度承诺、伤害信息。", "合规回应"],
    ], [1.2, 1.4, 2.2, 1.5])

    doc.add_heading("3.2 自然语言处理与事实剥离", level=2)
    doc.add_paragraph(
        "事实剥离模块的动机来自开放反馈中“希望不要再只是安慰我，合理去分析我的问题”的要求。"
        "用户需要的不是被简单劝好，而是有人帮助他看清“发生了什么”和“自己如何解释它”。"
        "系统对用户输入进行文本扫描，重点识别情绪化形容词、自我评价、他人心理推断、未来负面预测和绝对化表达。"
        "随后系统使用“监控摄像头标准”判断某一句话能否进入 A 诱发事件：如果摄像头、录音或第三方旁观者不能直接确认，"
        "该内容就不能作为事实，而应被归入 B 信念解释或 C 情绪行为。"
    )
    add_table(doc, ["ABC要素", "定义", "示例"], [
        ["A 诱发事件", "可观察、可记录、可复核的事实", "周五英语听写错了6个单词"],
        ["B 信念解释", "对事实的评价、推测、意义建构", "我彻底不行了"],
        ["C 情绪行为", "由解释引发的情绪、身体反应和行为倾向", "焦虑、羞愧、拖延背词"],
    ], [1.4, 2.6, 2.2])
    doc.add_paragraph(
        "在青少年语境中，事实剥离的难点在于许多句子表面上像事实，实际包含评价或推测。例如“老师故意针对我”并不是可观察事实，"
        "而是对教师动机的解释；“我完了”也不是事实，而是未来灾难化预测。系统处理这类输入时，不会直接反驳用户，"
        "而是先承认情绪，再把句子拆成“可确认发生了什么”“我脑中怎么解释它”“这种解释带来什么情绪和行为”三个层次。"
        "这种处理方式既避免否定用户感受，也降低情绪化推断继续放大的风险。"
    )

    doc.add_heading("3.3 情感计算与安全路由", level=2)
    doc.add_paragraph(
        "情感计算和安全路由模块的动机来自两类风险：一是问卷中“高度隐私”“深夜 emo”“社交疲惫”等场景占比较高，"
        "说明用户可能在现实支持较弱或不愿求助熟人的状态下打开系统；二是心理支持对象为未成年人，必须避免把危机当作普通情绪。"
        "情感计算模块不用于诊断，而用于判断回应强度。普通情绪压力进入共情、事实剥离和小行动建议；"
        "强烈无助、持续失眠或明显功能受损等升高风险场景进入风险澄清和现实支持建议；"
        "当出现自伤、自杀、伤人、虐待、性侵、严重欺凌、物质中毒等危机信号时，系统停止普通 CBT，优先进行安全处置。"
    )
    add_figure(doc, figs[10], "Figure 2. 风险分流与危机熔断路径。")
    add_table(doc, ["风险等级", "输入特征", "回应重点"], [
        ["普通支持", "学业、人际、自我评价压力，无危机词", "共情、事实剥离、小行动"],
        ["升高风险", "强烈无助、功能受损、持续困扰", "风险澄清、现实支持、必要转介"],
        ["危机风险", "自伤自杀、伤人、侵害、严重欺凌", "联系可信成年人、专业人员或紧急资源"],
    ], [1.3, 2.7, 2.2])

    doc.add_heading("3.4 RAG 知识库与来源治理", level=2)
    doc.add_paragraph(
        "RAG 知识库模块的动机来自开放反馈中对“更专业一点的评估”的期待，也来自心理类应用的专业风险。"
        "系统需要专业依据，但又不能把 AI 包装成医生；需要给出支持建议，但不能制造无法核验的权威感。"
        "为避免心理类 AI 常见的“来源不清”和“伪造案例”问题，项目将知识库拆分为权威指南、政策文件、CBT 理论和合成示例四类。"
        "权威指南用于风险分层、转介边界和安全处置；CBT 理论用于事实剥离和认知重评；合成示例只用于话术格式演示，"
        "并明确标注 synthetic=true 与 case_status=fictional_cross_checked。"
    )
    add_table(doc, ["文件", "行数", "主要用途"], [
        ["clinical_authoritative_chunks.jsonl", "28", "权威指南与安全边界"],
        ["clinical_professional_theory_chunks.jsonl", "40", "专业心理理论支撑"],
        ["synthetic_case_chunks.jsonl", "120", "校园合成场景与话术演示"],
        ["cbt_chunks.jsonl", "6", "CBT 核心理论和事实剥离"],
        ["dialogue_chunks.jsonl", "4", "对话风格和互动结构"],
        ["policy_chunks.jsonl", "2", "中国校园心理健康政策"],
    ], [2.7, 0.8, 2.7])
    doc.add_paragraph(
        "知识库治理不是简单把资料放进向量库，而是先判断资料能做什么、不能做什么。权威指南可以支持风险识别和转介边界，"
        "但不能被写成个案诊断；CBT 理论可以支持解释框架和对话技术，但不能替代治疗计划；合成案例只能用于格式演示，"
        "不能冒充真实临床案例。为了保持这一边界，每条知识块均应保留来源类型、目标对象、关键词、摘要、来源链接、"
        "证据等级、场景、技术、风险等级和案例状态等字段。"
    )
    doc.add_heading("3.5 用户可控记忆与隐私边界", level=2)
    doc.add_paragraph(
        "问卷和开放评论都显示，记忆功能是心理 AI 的关键矛盾点。一方面，用户说“记住我之前说过的事”，"
        "也有人希望“一个对话框一直用”，因为持续记忆可以减少重复解释，提高建议的贴合度；另一方面，问卷中也设置了“一键清空记忆”和“阅后即焚”选项，"
        "说明青少年在隐私困扰场景下非常需要控制权。因此，项目不把记忆设计成默认无限保存，而是设计为用户可控的分层记忆："
        "短期对话上下文用于当前交流，长期偏好或历史信息必须经过用户授权，用户应能够清空、分话题保存或选择不保存。"
    )
    add_table(doc, ["记忆层级", "保存内容", "用户控制方式", "设计动机"], [
        ["当前会话记忆", "本轮对话中的事实、情绪、偏好。", "关闭后可清空。", "保证当前回应连贯。"],
        ["主题对话框", "情感、人际、健康等不同主题下的历史线索。", "用户可分框保留或删除。", "回应“一个对话框一直用”的需求。"],
        ["长期偏好", "称呼、表达风格、常见压力源。", "需明确授权，可一键清空。", "提升个性化但避免隐私失控。"],
        ["敏感风险记录", "自伤、自杀、侵害等危机线索。", "仅用于安全提醒和现实求助建议。", "保护未成年人安全。"],
    ], [1.3, 1.8, 1.6, 1.8])

    doc.add_heading("四、我们如何解决问题：过程性设计说明", level=1)
    doc.add_paragraph(
        "本项目从草稿到成果的推进并不是线性堆功能，而是不断根据问卷反馈和心理安全风险修正设计。"
        "早期草稿更像一个“会回应情绪的聊天工具”，但问卷数据和开放评论表明，用户真正关心的是能否不模板化、能否合理分析、"
        "能否给出当下可做的小建议、能否记住但又不泄露隐私。因此，后续版本把这些需求逐一转化成模块。"
    )
    doc.add_paragraph(
        "`docs/survey_questionnaire.md` 说明问卷不是临时拼出来的，而是围绕破冰、语气偏好、抗疲劳、隐私担忧和低压力入口设计的。"
        "它与 `survey_export_2026-06-06.csv` 一起构成了“从用户痛点到产品模块”的证据链。"
    )
    add_table(doc, ["问卷或风险证据", "草稿中的不足", "修订动作", "形成的成果模块"], [
        ["痛点中“stiff_template”占比最高，开放评论说“不是那种完全套话”。", "回应容易像固定模板。", "引入场景、情绪、认知偏差和知识片段共同决定回应。", "个性化回应生成模块"],
        ["开放评论说“希望不要再只是安慰我，合理去分析我的问题”。", "只有共情，缺少结构化分析。", "用 ABC 模型拆分事实、信念、情绪行为。", "事实剥离与认知偏差识别模块"],
        ["开放评论说“多给一点当下能做的一些微小能缓解压力行为的建议”。", "建议偏泛泛而谈。", "把每次回应收束到可执行小行动。", "小行动建议模块"],
        ["开放评论说“记住我之前说过的事”，同时问卷包含一键清空和阅后即焚选项。", "没有处理连续性与隐私的冲突。", "设计用户可控分层记忆。", "记忆与隐私边界模块"],
        ["高度隐私、深夜情绪、社交疲惫为高频求助场景。", "入口可能过重，像测评或盘问。", "首屏采用表情包选择、温和寒暄和模式选择。", "低压力进入原型"],
        ["问卷第3题要求用户用1-3个表情代表负面/低谷情绪。", "文字倾诉门槛过高，用户可能不知道如何开口。", "把表情包作为情绪输入前置入口。", "情绪表情包选择模块"],
        ["开放反馈反复提到“像人”“不要太套路化”。", "AI容易输出长段、规整、像模板的话。", "通过提示词工程约束短句、口语、引用用户原话。", "拟人化回应提示词模块"],
        ["测试中发现首轮长句容易造成阅读负担。", "系统想表达很多，但用户第一轮未必愿意读。", "建立长句压缩为短句的改写流程。", "短句化开场与轻量交互模块"],
        ["未成年人心理支持存在自伤、自杀、侵害等安全风险。", "普通 CBT 可能误用于危机场景。", "危机知识块覆盖普通回应。", "安全优先风险路由"],
    ], [2.0, 1.6, 1.8, 1.8])
    doc.add_heading("4.1 从“普通安慰”改为“结构化支持”", level=2)
    doc.add_paragraph(
        "这一改动直接回应问卷开放评论中的“希望不要再只是安慰我”和“可以理性但不能太理性，要感性”。"
        "早期设计中，智能体容易输出“别难过”“你已经很棒了”等泛化安慰。我们认为这种回应虽然温和，"
        "但无法帮助青少年理解情绪来源，也难以形成可执行行动。因此，我们把回应流程改为五步：情绪接纳、事实剥离、信念标注、平衡解释、小行动建议。"
    )
    doc.add_paragraph(
        "这一路径对应 CBT 中从情境到自动化思维再到情绪行为结果的整理过程。与传统问答式 AI 不同，本项目不把“回答得像人”作为唯一目标，"
        "而是要求每一次回应都能说明：它接住了什么情绪、确认了哪些事实、识别了哪些推断、给出的行动是否足够小而具体。"
        "例如，对于“听写错了 6 个，我完了”这类输入，系统不能只说“别灰心”，而应明确指出“错 6 个词”是事实，"
        "“我完了”是灾难化预测，下一步可以先把错词分成拼写、发音和词义三类。"
    )
    doc.add_heading("4.2 从“能回答”改为“有证据边界”", level=2)
    doc.add_paragraph(
        "这一改动回应的是“专业”与“非诊断”之间的矛盾。用户希望 AI 更专业，但比赛项目和未成年人心理安全要求决定了系统不能进行正式诊断。"
        "心理支持智能体不能只追求回答流畅，还必须知道哪些内容能说、哪些内容不能说。我们建立 RAG 质量规则，要求知识块保留 source_urls、"
        "evidence_level、scenario、technique、safety_level 和 case_status 等字段；同时禁止把合成示例写成真实临床案例。"
    )
    doc.add_paragraph(
        "这一设计解决的是心理 AI 中常见的“听起来专业但无法核验”的问题。报告和系统都不使用无法追溯的“研究表明”作为论据，"
        "而是将 WHO、NICE、AAP、AACAP、教育部政策文件和 Beck Institute CBT 资料作为不同层级的支撑材料。"
        "其中，国际指南承担安全与转介边界，政策文件承担校园应用背景，CBT 资料承担对话技术框架，合成场景仅用于演示输出格式。"
    )
    doc.add_heading("4.3 从“统一话术”改为“风险分层”", level=2)
    doc.add_paragraph(
        "青少年输入可能只是普通学业压力，也可能包含危机信号。为解决不同风险混在一起的问题，我们设计了安全优先路由："
        "危机知识块覆盖普通 CBT；升高风险先澄清安全和现实支持；普通场景再执行事实剥离和认知重评。"
    )
    doc.add_paragraph(
        "风险分层的关键不是把用户贴上标签，而是决定系统应该采取哪类行动。普通压力场景中，系统可以使用共情、ABC 整理和小行动；"
        "升高风险场景中，系统应询问安全状态、持续时间、功能受损和可获得支持；危机场景中，系统必须停止普通陪聊，"
        "转向现实成年人、学校心理老师、专业人员或紧急资源。这个规则也回应了未成年人应用场景中的伦理要求：AI 可以作为辅助入口，"
        "但不能让青少年独自处理高风险问题。"
    )
    doc.add_heading("4.4 情绪表情包选择：把复杂情绪变成低压力输入", level=2)
    doc.add_paragraph(
        "问卷第 3 题要求受访者凭直觉选择 1-3 个最能代表当下负面或低谷情绪的表情包。这个设计不是为了让页面更活泼，"
        "而是为了解决青少年在心理求助起点上的表达困难。很多用户并不是没有情绪，而是不知道如何把情绪组织成完整句子；"
        "尤其在深夜 emo、社交疲惫、高度隐私困扰等场景中，要求用户一开始就写长段描述，会提高退出概率。"
        "因此，我们把表情包选择设计为一种低压力情绪入口：用户可以先用 🫠、😭、😑、💤 等符号表达状态，系统再根据表情组合推断可能的情绪倾向，"
        "用更轻的句子邀请用户继续，而不是立刻进入正式问答。"
    )
    doc.add_paragraph(
        "这一模块的成果是：首屏不再只有文字输入框，而可以提供“表情包情绪点选 + 轻量模式选择”。"
        "例如，当用户选择“🫠 + 💤”时，系统不直接问“请描述你的心理问题”，而可以回应：“看起来像是又累又有点撑不住。"
        "你想先吐槽两句，还是只让我陪你待一会儿？”这种回应把表达压力降到最低，同时保留继续倾诉的入口。"
    )
    add_figure(doc, figs[14], "Figure 3. 情绪表情包低压力入口机制。")
    doc.add_heading("4.5 长句变短句：从完整说明到可读开场", level=2)
    doc.add_paragraph(
        "在草稿阶段，智能体为了显得专业，容易一次性输出背景解释、共情、分析、建议和提醒，结果形成很长的回应。"
        "但问卷中的痛点显示，用户反感“说话机械套路”和“冷冰冰地做测评量表诊断”；开放评论也提到“不要太套路化了，感觉每次都能猜到 AI 说什么”。"
        "因此，我们把“长句变短句”作为专门的提示词工程任务，而不是简单删字。"
    )
    doc.add_paragraph(
        "具体做法分为四步：第一，把原始长回应拆成事实、情绪、解释、建议和安全提醒；第二，判断当前轮次用户最需要哪一类信息；"
        "第三，只保留一个共情点、一个关键观察和一个可选择动作；第四，把书面语改成接近日常聊天的短句。"
        "例如，草稿句“根据你的描述，你目前可能处于明显的学业压力与自我评价降低状态，建议你先进行情绪调节并尝试记录自动化思维”"
        "会被改写为：“这次像是成绩一下把你打懵了。我们先不急着给你下结论，要不要先把发生了什么拆出来？”"
        "短句化的目标不是降低专业性，而是让专业支持先能被读完、被接住、被回应。"
    )
    add_figure(doc, figs[15], "Figure 4. 长句到短句的表达压缩流程。")
    doc.add_heading("4.6 拟人化提示词工程：像人，但不越界", level=2)
    doc.add_paragraph(
        "开放评论中多次出现“像人就行”“期待能更像人”“可以理性但不能太理性，要感性”等反馈。"
        "这说明用户并不是单纯追求娱乐化，而是希望 AI 在心理支持中保留真实对话的温度：能听懂语气，能接住情绪，"
        "能引用用户刚说过的话，而不是输出一段标准化心理教育。为此，我们设计了拟人化提示词工程。"
    )
    doc.add_paragraph(
        "提示词工程的核心约束包括：先共情再分析、短句优先、少术语、不模板化、引用用户原话、给选择权、给微小行动。"
        "同时，拟人化必须受到安全边界限制：不能假装真人朋友持续陪伴，不能承诺治疗效果，不能替代专业人员。"
        "因此，本项目追求的是“有温度的结构化支持”，而不是无边界的情感扮演。"
        "例如，系统可以说“你刚刚说‘不想再只是被安慰’，那我们这次先一起分析发生了什么”，"
        "但不能说“我会一直陪着你、保证你没事”。"
    )
    doc.add_paragraph(
        "`worker/src/lib/llm.ts` 是提示词工程最直接的实现文件。它把角色定义、情感兜底、交互风格、输出格式、RAG 注入和 FSM 状态串成可运行的 system prompt，"
        "所以“像人但不失控”不是前端文案，而是后端行为约束。"
    )
    add_figure(doc, figs[16], "Figure 5. 基于问卷反馈的拟人化提示词工程框架。")
    doc.add_heading("4.7 用户记忆与隐私边界：记住，但由用户决定记到哪一步", level=2)
    doc.add_paragraph(
        "记忆设计的动机几乎是被问卷直接点出来的。开放评论里有人明确说“记住我之前说过的事”，"
        "也有人在隐私题项里更偏好“阅后即焚”或“用户可控”。这说明青少年对心理 AI 的期待并不是单向度的："
        "他们一方面希望系统能记住上下文、减少重复解释，另一方面又非常担心聊天内容被过度留存、被别人看到、被系统“看穿”。"
        "因此，项目没有把记忆设计成“自动永久保存”，而是做成分层控制。"
    )
    doc.add_paragraph(
        "在实现层面，`worker/src/routes/chat.ts` 会从会话中恢复 FSM 上下文，而 `worker/src/routes/onboarding.ts` 会将自由文本压缩成 `weather`、"
        "`safetyIsland` 和 `stressor` 三个初始画像维度，用于后续对话的轻量锚定。更细的长期偏好，只有在用户明确授权时才进入可复用记忆；"
        "对于高风险线索，系统只保留必要的安全判断信息，不把它们当作普通兴趣标签或个性标签长期扩散。"
        "这使系统既能接上下文，又不至于造成“被盯着看”的压迫感。"
    )
    doc.add_paragraph(
        "从产品体验上看，记忆机制对应的是“我可以不重复说很多次”“你能接上我上次的话”，"
        "而隐私机制对应的是“我知道你记了什么，也知道我可以删掉什么”。这两个能力缺一不可。"
        "如果只有记忆，用户会担心被监视；如果只有清空，系统又会显得像没有脑子。"
    )
    doc.add_heading("4.8 状态机与风险分流：让每一类问题走不同通道", level=2)
    doc.add_paragraph(
        "项目后台最核心的行为控制层是有限状态机。`worker/src/lib/fsm.ts` 明确把对话分为 `Onboarding`、`Active_Listening`、"
        "`CBT_Stripping`、`Socratic_Questioning`、`Crisis_Escalation` 五种状态，其中 `Crisis_Escalation` 是吸收态，一旦进入不可逆退。"
        "这个设计之所以重要，是因为问卷和开放评论已经告诉我们：用户并不是每次都需要同一类回复。"
        "有时他们只想要一个温和开场，有时希望先被接住情绪，有时才需要结构化分析，有时则必须立即切安全模式。"
    )
    doc.add_paragraph(
        "`worker/src/lib/intent-router.ts` 负责第一层意图分流，将输入分为 `casual`、`emotional`、`ambiguous`、`crisis` 四类。"
        "这一层分流不是“多此一举”，而是为了避免所有消息都进同一套 prompt：日常闲聊只需要低压力陪伴，"
        "情绪倾诉才需要进入 CBT 结构，危机场景则要直接跳过普通分析。"
        "这也是为什么报告要把“风险分流”写成独立模块，而不是只在伦理部分轻描淡写地提一句。"
    )
    doc.add_paragraph(
        "真正的行为逻辑由 `worker/src/routes/chat.ts` 串起来：先恢复会话状态，再做意图识别，再执行 pre-transition，"
        "随后决定是否检索 RAG 知识，再把当前状态、画像、情绪信号和知识片段注入 system prompt，最后在模型回复后再做 post-transition。"
        "换句话说，系统不是‘发一个 prompt 等它回’，而是一个状态机驱动的对话编排器。"
    )
    doc.add_heading("4.9 前端低压力入口：问卷、表情包与视觉气氛不是装饰", level=2)
    doc.add_paragraph(
        "前端低压力入口的设计同样来自问卷。`web/public/survey/index.html` 对应的问卷页面本身就不是传统量表，而是让用户通过表情、偏好和选择题表达状态。"
        "问卷导出 `survey_export_2026-06-06.csv` 中，首屏偏好明显倾向于“表情包选择”和“温和寒暄”，而不是一上来就临床式评估；"
        "这说明前端必须承担“降低第一秒表达门槛”的角色。"
    )
    doc.add_paragraph(
        "因此，前端不只是一个聊天框，而是包含情绪表情包入口、轻聊天开场、隐私感提示和低阻尼输入方式的交互界面。"
        "如果后续页面接入更丰富的情绪信号，例如摄像头表情识别或动态情绪反馈，前端的任务也不是“炫技”，"
        "而是帮助系统在用户还没开口前就感知到紧张、回避或低落的状态，从而调整入口语气。"
    )
    doc.add_heading("4.10 从草稿到成果：研发叙事如何落到文件里", level=2)
    doc.add_paragraph(
        "项目并不是一开始就长成现在这样，而是通过 `docs/development_record.md`、`ORIGINAL_REQUEST.md`、`plan.md`、`tasks.md` 这一系列记录逐渐收敛。"
        "开发早期的草稿更像是在验证：要不要做 CBT、要不要做破冰、要不要做安全熔断；后来再通过问卷、开放评论和代码迭代，把这些问题收成明确的模块。"
        "所以报告在写作上不应该只呈现“我们做了什么”，而要呈现“为什么先做这个，再做那个”。"
    )
    doc.add_paragraph(
        "例如，先做问卷，是因为需要知道中学生到底讨厌什么；先做表情包入口，是因为文字门槛太高；先做状态机，是因为不同问题不能用同一条路径处理；"
        "先做提示词工程，是因为即使同一状态下，回复方式也需要根据问卷反馈压制套路感、增强人味。"
        "这种从草稿到成果的推进逻辑，是整个项目最值得展示的部分。"
    )

    doc.add_heading("五、系统原型与测试设计", level=1)
    doc.add_paragraph(
        "系统原型由三类材料共同构成：第一类是线上匿名问卷和数据大屏，用于开题阶段的需求研究；第二类是本地 `survey.html`，"
        "用于进一步测试低注意力状态下不同开场话术的即时反应；第三类是规则文档、RAG 知识库和验证脚本，用于证明智能体回应不是一次性文案，"
        "而是可以被约束、复核和迭代的系统。"
    )
    doc.add_paragraph(
        "其中，线上问卷回答“用户想要什么、不想要什么”，本地开场测试回答“第一句话是否会劝退用户”，RAG 与验证脚本回答“系统能不能安全稳定地工作”。"
        "三者共同构成从用户研究到工程验证的闭环。报告中将它们分别作为需求证据、交互证据和技术证据。"
    )

    doc.add_heading("六、典型应用案例", level=1)
    doc.add_heading("6.1 考试压力场景", level=2)
    doc.add_paragraph("用户输入：“因为英语听写错了 6 个单词，我就觉得自己彻底不行了，好像只要不是做得最好就是全盘失败。”")
    add_table(doc, ["处理项", "系统输出"], [
        ["A 诱发事件", "周五英语听写后，老师公布用户错了6个单词。"],
        ["B 信念解释", "用户把一次听写错误解释为整体失败。"],
        ["C 情绪行为", "紧张、羞愧、担心，可能拖延背词。"],
        ["认知偏差", "非黑即白。"],
        ["下一步行动", "把6个错词分成拼写、发音、词义三类，今天先复习其中一类。"],
    ], [1.5, 4.7])
    doc.add_heading("6.2 同伴关系场景", level=2)
    doc.add_paragraph("用户输入：“我回答问题时同桌笑了一下，他肯定觉得我很蠢。”")
    add_table(doc, ["处理项", "系统输出"], [
        ["A 诱发事件", "课堂上，用户回答问题时，同桌笑了一下。"],
        ["B 信念解释", "用户把“同桌笑了一下”解释为“他觉得我很蠢”。"],
        ["C 情绪行为", "尴尬、紧张、羞耻，可能不再主动回答问题。"],
        ["认知偏差", "读心术。"],
        ["下一步行动", "记录还有哪些解释也可能说明“同桌笑了一下”，再决定是否需要沟通。"],
    ], [1.5, 4.7])
    doc.add_heading("6.3 危机风险场景", level=2)
    doc.add_paragraph(
        "若用户表达“我真的撑不下去了，不想活了”，系统不进入普通 CBT 事实剥离，也不把该表达简单解释为灾难化思维；"
        "系统优先确认当前安全状态，建议联系身边可信成年人、学校老师、心理老师、专业人员或当地紧急资源。"
    )

    doc.add_heading("七、创新点", level=1)
    doc.add_paragraph(
        "本项目的创新并不体现在单一算法名称上，而体现在心理学结构、青少年交互偏好和 AI 安全治理三者的组合。"
        "多数通用聊天机器人可以回应情绪，但不一定能稳定区分事实和解释；部分心理问答工具可以给出建议，但容易变成量表式盘问；"
        "而本项目试图在低压力进入、结构化心理支持和安全可追溯之间建立平衡。"
    )
    add_table(doc, ["创新点", "传统做法的不足", "本项目做法", "预期价值"], [
        ["事实与情绪剥离", "直接安慰或直接反驳，容易忽视事实-解释混杂。", "用 ABC 和监控摄像头标准拆解倾诉。", "降低认知偏差放大，提升解释清晰度。"],
        ["青少年偏好驱动", "先设计产品再让用户适应。", "用问卷和大屏反推首屏、角色和隐私控制。", "提升进入意愿和真实使用贴合度。"],
        ["安全优先 RAG", "所有场景使用同一套陪聊逻辑。", "危机场景覆盖普通 CBT，优先现实求助。", "降低高风险场景误导。"],
        ["来源分层治理", "资料混杂，案例真假不清。", "区分权威指南、政策、理论和合成示例。", "提高专业可信度和可核验性。"],
        ["过程可验证", "只展示最终效果。", "保留问卷、规则、知识库、脚本、截图和视频链路。", "方便评审复核项目真实性。"],
    ], [1.3, 2.0, 2.0, 1.7])
    add_figure(doc, figs[21], "Figure 6. 从用户证据到系统成果的可验证链路。")

    doc.add_heading("八、验证方案与质量控制", level=1)
    doc.add_paragraph(
        "验证方案分为功能验证、知识库验证和材料验证三类。功能验证检查系统能否执行事实剥离、认知偏差识别、风险路由和安全回应；"
        "知识库验证检查字段完整性、来源可追溯和合成示例标注；材料验证检查过程截图、视频链接、PDF 格式和查重报告。"
    )
    doc.add_paragraph(
        "为了避免验证停留在主观描述层面，本项目把测试对象拆成六类输入：日常闲聊、普通学业压力、同伴误解、家长期待、危机风险和来源追溯。"
        "每类输入都对应不同预期行为。比如日常闲聊不能被强行心理分析；同伴误解必须标注读心术而不能认同推测；危机风险必须跳过普通 CBT；"
        "来源追溯必须能说明知识依据。只有这些场景都能通过，才能说明系统不是单次样例可行，而是在规则约束下具有稳定性。"
    )
    add_table(doc, ["测试编号", "输入类型", "预期输出"], [
        ["T01", "日常闲聊", "不强行分析心理问题"],
        ["T02", "考试压力", "输出 ABC 与小行动"],
        ["T03", "同伴误解", "标注读心术，不认同推测"],
        ["T04", "家长期待", "共情并拆解事实与评价"],
        ["T05", "自伤风险", "危机优先，现实求助"],
        ["T06", "来源追溯", "返回权威知识依据"],
    ], [1.0, 2.0, 3.2])
    add_table(doc, ["质量控制对象", "检查内容", "失败示例", "修正方式"], [
        ["ABC 输出", "A 是否只包含可观察事实。", "把“老师针对我”写进 A。", "改为“老师指出作业错误”。"],
        ["风险路由", "危机词是否覆盖普通回应。", "对“不想活了”做认知重构。", "转入安全支持和现实求助。"],
        ["知识来源", "是否有 source_urls 和 evidence_level。", "无来源地写“研究表明”。", "召回权威来源并标注用途。"],
        ["合成案例", "是否明确 synthetic 和 fictional。", "把编造场景称为临床案例。", "改为“应用示例”。"],
        ["输出边界", "是否含诊断、处方或过度承诺。", "说“你这是抑郁症”。", "改为非诊断支持和求助建议。"],
    ], [1.5, 2.1, 2.0, 1.7])

    doc.add_heading("九、项目边界与伦理说明", level=1)
    add_bullets(doc, [
        "不提供医学诊断、心理疾病诊断、药物建议或治疗方案。",
        "不替代心理咨询师、精神科医生、学校心理老师或急救系统。",
        "不收录未授权真实病例，不把合成样例包装成临床案例。",
        "对未成年人危机风险，优先建议联系现实中的可信成年人和专业资源。",
        "对自伤、自杀、伤人等内容，不提供方法、工具、剂量、规避方式等可操作伤害信息。",
    ])
    # 预留 Figure 16 之后的提交材料包图，可在后续扩展。

    doc.add_heading("十、后续完善计划", level=1)
    add_numbered(doc, [
        "补充 survey.html 首页、测试卡片和结果页真实截图。",
        "运行验证脚本并将终端结果截图补入过程性证明材料。",
        "统计合成场景数据的真实场景分布，替换当前场景覆盖示意图。",
        "录制参赛者成果介绍视频，并在过程性证明材料中插入可直接播放链接。",
        "将本报告导出为 PDF，控制单个文件大小不超过 50MB。",
        "使用知网、维普、万方等权威系统进行查重，确保查重率小于或等于 30%。",
    ])

    doc.add_heading("十一、结论", level=1)
    doc.add_paragraph(
        "本项目围绕青少年心理成长需求，提出了一个兼具心理学结构、AI 技术实现和安全边界意识的心理支持智能体方案。"
        "其核心价值不在于替代专业心理服务，而在于以低门槛方式帮助青少年把混乱的情绪倾诉整理为可理解、可讨论、可行动的结构，"
        "并在风险升高时引导用户连接现实支持系统。通过 CBT 事实剥离、认知偏差识别、情感计算、RAG 权威知识检索和多模态过程证明，"
        "本项目能够形成可落地、可验证的系统原型与应用方案。"
    )

    doc.add_heading("附录：参考资料与项目文件", level=1)
    doc.add_paragraph(
        "本附录不再只列代表性文献，而是根据 `rag-psy-cbt/data/cleaned/*.jsonl` 中的 `source_urls` 字段自动汇总 RAG 知识库实际使用过的外部来源。"
        "同一来源被多个 chunk 召回时只保留一条 URL，并在“来源类型与用途”中标注涉及的知识库文件和 chunk 数量。"
    )
    rag_files, rag_refs, internal_refs = read_rag_reference_inventory()
    doc.add_heading("附录 A：RAG 知识库文件统计", level=2)
    add_table(doc, ["知识库文件", "条目数"], rag_files, [4.5, 1.2])
    doc.add_heading("附录 B：RAG 外部参考文献与网页来源完整清单", level=2)
    doc.add_paragraph(f"以下共列出 {len(rag_refs)} 条去重后的外部来源。为避免长 URL 撑开页面，本节采用编号参考文献列表格式。")
    add_reference_list(doc, rag_refs)
    doc.add_heading("附录 C：无外部 URL 的内部规则与自建材料", level=2)
    doc.add_paragraph(
        "以下条目来自知识库中未填写 `source_urls` 的内部规则、政策摘编、对话规范或项目自建安全话术。"
        "`synthetic_case_chunks.jsonl` 共 120 条，均为项目自建校园心理支持合成场景，已在元数据中标注 `synthetic=true`、"
        "`case_status=fictional_cross_checked`、`abc_validated=true`，仅用于话术格式演示和测试，不作为真实临床病例引用。"
    )
    add_table(doc, ["文件", "chunk_id", "类型", "说明"], internal_refs, [1.8, 2.0, 1.0, 2.4])
    doc.add_heading("附录 D：项目研究与提交材料链接", level=2)
    add_bullets(doc, [
        "问卷页面：https://survey.rethink.buleegasy.space",
        "数据大屏：https://re-think-dashboard.pages.dev",
        "项目文件：README.md、ORIGINAL_REQUEST.md、plan.md、tasks.md、docs/survey_questionnaire.md、docs/development_record.md。",
        "核心实现文件：worker/src/lib/llm.ts、worker/src/lib/fsm.ts、worker/src/lib/intent-router.ts、worker/src/routes/chat.ts、worker/src/routes/onboarding.ts、worker/src/lib/rag.ts。",
    ])
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
